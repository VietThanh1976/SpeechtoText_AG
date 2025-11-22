import streamlit as st
import speech_recognition as sr
import io
import os
import json
import time

# Thư viện cho file DOCX
from docx import Document 
# Thư viện xử lý audio file (thay thế pydub)
import librosa 
import soundfile as sf 
# Thư viện ghi âm từ micro trình duyệt (thay thế pyaudio)
from streamlit_mic_recorder import mic_recorder 
# Thư viện mã hóa (Tùy chọn, dùng để kiểm tra mật khẩu an toàn hơn)
from st_hashing import Hashing 

# =========================================================================
# I. KHỞI TẠO VÀ CẤU HÌNH BAN ĐẦU
# =========================================================================

r = sr.Recognizer()
h = Hashing()

# Khởi tạo Session State
if 'audio_buffer' not in st.session_state:
    st.session_state.audio_buffer = None
if 'last_transcription_text' not in st.session_state:
    st.session_state.last_transcription_text = ""
if 'last_audio_data' not in st.session_state:
    st.session_state.last_audio_data = None
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if 'user_name' not in st.session_state:
    st.session_state.user_name = None

# Tải thông tin người dùng từ file JSON
try:
    with open("user_credentials.json", "r") as f:
        USERS = json.load(f)
except FileNotFoundError:
    st.error("Lỗi: Không tìm thấy file 'user_credentials.json'. Vui lòng kiểm tra lại cấu trúc file.")
    st.stop()
except Exception as e:
    st.error(f"Lỗi khi đọc file JSON: {e}")
    st.stop()


# =========================================================================
# II. CÁC HÀM HỖ TRỢ
# =========================================================================

def transcribe_audio_from_file_path(file_path):
    """Sử dụng SpeechRecognition để chuyển đổi file WAV thành văn bản."""
    r = sr.Recognizer()
    try:
        with sr.AudioFile(file_path) as source:
            audio = r.record(source) 
        text = r.recognize_google(audio, language="vi-VN")
        return text
    except sr.UnknownValueError:
        return "Không thể nhận dạng giọng nói từ tệp âm thanh này."
    except sr.RequestError as e:
        return f"Lỗi kết nối hoặc API: {e}"
    except Exception as e:
        return f"Lỗi xử lý tệp: {e}"

def process_uploaded_file(uploaded_file):
    """Xử lý và chuyển đổi file đã tải lên."""
    st.session_state.last_transcription_text = ""
    # Code xử lý file dùng librosa và soundfile (giống như trước)
    # ... (Giữ nguyên logic xử lý file của bản hoàn chỉnh trước đó)
    temp_input_path = "temp_input_audio" + os.path.splitext(uploaded_file.name)[1]
    temp_wav_path = "temp_converted_audio.wav"
    
    try:
        with open(temp_input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        y, sr_librosa = librosa.load(temp_input_path, sr=None) 
        sf.write(temp_wav_path, y, sr_librosa)
        
        st.info("Đang nhận dạng giọng nói...")
        result_text = transcribe_audio_from_file_path(temp_wav_path)
        
        st.session_state.last_transcription_text = result_text

    except Exception as e:
        st.session_state.last_transcription_text = f"Lỗi xử lý tệp: {e}"
    finally:
        if os.path.exists(temp_input_path):
            os.remove(temp_input_path)
        if os.path.exists(temp_wav_path):
            os.remove(temp_wav_path)


def create_docx(text, filename="transcribed_document.docx"):
    """Tạo một file DOCX từ văn bản đã chuyển đổi."""
    document = Document()
    document.add_heading('Văn bản đã chuyển đổi', 0)
    document.add_paragraph(text)

    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io.read(), filename


def check_login():
    """Hiển thị form đăng nhập và xử lý xác thực."""
    
    if st.session_state.logged_in:
        return True

    # Tạo giao diện đăng nhập ở Sidebar
    with st.sidebar.form("login_form"):
        st.sidebar.title("Đăng nhập 🔑")
        username = st.text_input("Tên đăng nhập")
        password = st.text_input("Mật khẩu", type="password")
        login_button = st.form_submit_button("Đăng nhập")
        
        if login_button:
            if username in USERS:
                hashed_password = USERS[username]["password"]
                
                # Kiểm tra mật khẩu mã hóa
                if h.check_hash(password, hashed_password):
                    st.session_state.logged_in = True
                    st.session_state.user_name = USERS[username]["name"]
                    st.success(f"Chào mừng, {st.session_state.user_name}!")
                    time.sleep(1) # Tạm dừng để người dùng thấy thông báo
                    st.rerun() 
                else:
                    st.error("Mật khẩu không đúng!")
            else:
                st.error("Tên đăng nhập không tồn tại!")
    
    return False

def logout():
    """Xử lý đăng xuất."""
    st.session_state.logged_in = False
    st.session_state.user_name = None
    st.session_state.last_transcription_text = "" # Xóa kết quả cũ
    st.sidebar.success("Đã đăng xuất.")
    st.rerun()

# =========================================================================
# III. ỨNG DỤNG CHÍNH (MAIN APP)
# =========================================================================

# 1. Kiểm tra và xử lý Đăng nhập
if not check_login():
    st.title("Vui lòng Đăng nhập để sử dụng Ứng dụng")
    st.info("Sử dụng: **admin / 123456** hoặc **user1 / password**")
    st.stop()
    
# 2. Nếu đã đăng nhập, hiển thị nội dung chính
st.title("🎤 Ứng Dụng Chuyển Giọng Nói Thành Văn Bản")
st.markdown(f"**Người dùng:** **{st.session_state.user_name}** | **Tên đăng nhập:** `{st.session_state.user_name}`")

# Nút Đăng xuất
if st.sidebar.button("Đăng xuất"):
    logout()
    
# --- Chọn Phương thức ---
st.markdown("---") 

method = st.radio(
    "Chọn phương thức nhập liệu:",
    ('Tải lên File Âm thanh', 'Ghi âm trực tiếp từ Micro')
)

### PHƯƠNG THỨC 1: Tải lên File Âm thanh
if method == 'Tải lên File Âm thanh':
    uploaded_file = st.file_uploader(
        "Tải lên tệp âm thanh (.wav, .mp3, etc.):",
        type=['wav', 'mp3', 'ogg', 'flac']
    )
    if uploaded_file is not None:
        if st.button('🚀 Chuyển đổi File thành Văn bản'):
            with st.spinner('Đang tải và xử lý file...'):
                process_uploaded_file(uploaded_file)
            
### PHƯƠNG THỨC 2: Ghi âm trực tiếp từ Micro
elif method == 'Ghi âm trực tiếp từ Micro':
    st.subheader("🎙️ Ghi Âm Trực Tiếp")
    st.caption("Ghi âm bằng micro của trình duyệt.")

    # Widget ghi âm
    audio_data = mic_recorder(
        start_prompt="Bắt đầu Ghi Âm",
        stop_prompt="Dừng Ghi Âm",
        key='mic_recorder',
        format="wav"
    )

    if audio_data:
        st.session_state.audio_buffer = audio_data['bytes']
        st.session_state.last_audio_data = audio_data['bytes']
        st.audio(st.session_state.audio_buffer, format='audio/wav') 
        
        st.download_button(
            label="⬇️ Tải xuống File Âm thanh (.wav)",
            data=st.session_state.last_audio_data,
            file_name="ghi_am_mic.wav",
            mime="audio/wav"
        )
    
    if st.session_state.audio_buffer is not None:
        if st.button('✅ Chuyển đổi Giọng nói'):
            st.session_state.last_transcription_text = ""
            
            temp_wav_path = "mic_recording_temp.wav"
            
            try:
                with open(temp_wav_path, "wb") as f:
                    f.write(st.session_state.audio_buffer)

                with st.spinner('Đang nhận dạng giọng nói...'):
                    result_text = transcribe_audio_from_file_path(temp_wav_path)
                
                st.session_state.last_transcription_text = result_text

            except Exception as e:
                st.session_state.last_transcription_text = f"Lỗi xử lý: {e}"
            finally:
                if os.path.exists(temp_wav_path):
                    os.remove(temp_wav_path)


# =========================================================================
# IV. HIỂN THỊ KẾT QUẢ VÀ TÙY CHỌN TẢI XUỐNG
# =========================================================================
if st.session_state.last_transcription_text:
    st.markdown("---")
    st.subheader("✅ Văn bản đã chuyển đổi:")
    
    st.text_area("Kết quả:", st.session_state.last_transcription_text, height=250)

    if "Không thể" not in st.session_state.last_transcription_text and "Lỗi" not in st.session_state.last_transcription_text:
        
        col1, col2 = st.columns(2)
        
        # Nút tải xuống file DOCX
        docx_bytes, docx_filename = create_docx(st.session_state.last_transcription_text)
        col1.download_button(
            label="💾 Tải xuống MS Word (.docx)",
            data=docx_bytes,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Nút tùy chọn tải xuống file TXT
        col2.download_button(
            label="📝 Tải xuống Văn bản thuần (.txt)",
            data=st.session_state.last_transcription_text.encode('utf-8'),
            file_name="transcribed_text.txt",
            mime="text/plain"
        )
